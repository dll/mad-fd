from docx import Document

# 打开DOCX文件
doc = Document('计科22《移动应用开发》课程达成评价表格-课程目标达成评价报告.docx')

# 提取所有段落文本
print("文档内容:")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"段落 {i}: {para.text}")

# 提取所有表格
print("\n表格内容:")
for i, table in enumerate(doc.tables):
    print(f"表格 {i}:")
    for row in table.rows:
        row_text = [cell.text.strip() for cell in row.cells]
        print(f"  {row_text}")