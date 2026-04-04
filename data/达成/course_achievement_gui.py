"""
课程达成度计算系统 - 图形界面版本
支持导入Excel/Word文件、可视化展示、生成报告、打印功能
可配置不同课程的权重，通用化设计
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, scrolledtext
import os
from datetime import datetime
import json
from pathlib import Path

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
plt.rcParams['axes.unicode_minus'] = False


class CourseAchievementGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("课程达成度计算系统")
        self.root.geometry("1200x800")
        
        # 数据存储
        self.students = []
        self.avg_achievements = {}
        self.weighted_achievement = 0
        self.course_config = {}
        self.excel_file_path = None
        self.syllabus_file_path = None
        self.output_dir = './output'
        
        # 默认课程目标权重
        self.objective_weights = {
            '课程目标1': 0.15,
            '课程目标2': 0.25,
            '课程目标3': 0.30,
            '课程目标4': 0.30
        }
        
        # 默认评估类型权重
        self.assessment_weights = {
            '平时': 0.20,
            '实验': 0.30,
            '期末': 0.50
        }
        
        # 创建界面
        self.create_widgets()
        
        # 创建输出目录
        os.makedirs(self.output_dir, exist_ok=True)
    
    def create_widgets(self):
        # 主框架
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 配置网格权重
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(1, weight=1)
        
        # 标题
        title_label = ttk.Label(
            main_frame, 
            text="课程达成度计算系统", 
            font=("Microsoft YaHei", 18, "bold")
        )
        title_label.grid(row=0, column=0, columnspan=2, pady=10)
        
        # 左侧面板 - 导入和配置
        left_panel = ttk.LabelFrame(main_frame, text="导入与配置", padding="10")
        left_panel.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5)
        
        # 文件导入部分
        import_frame = ttk.LabelFrame(left_panel, text="文件导入", padding="10")
        import_frame.pack(fill=tk.X, pady=5)
        
        # Excel文件导入
        ttk.Button(import_frame, text="导入Excel成绩文件", 
                 command=self.import_excel).pack(fill=tk.X, pady=2)
        self.excel_label = ttk.Label(import_frame, text="未选择文件", foreground="gray")
        self.excel_label.pack(fill=tk.X)
        
        # Word文件导入
        ttk.Button(import_frame, text="导入Word大纲文件", 
                 command=self.import_word).pack(fill=tk.X, pady=2)
        self.word_label = ttk.Label(import_frame, text="未选择文件", foreground="gray")
        self.word_label.pack(fill=tk.X)
        
        # 课程配置部分
        config_frame = ttk.LabelFrame(left_panel, text="课程配置", padding="10")
        config_frame.pack(fill=tk.X, pady=5)
        
        # 课程名称
        ttk.Label(config_frame, text="课程名称:").pack(anchor=tk.W)
        self.course_name_entry = ttk.Entry(config_frame)
        self.course_name_entry.pack(fill=tk.X, pady=2)
        self.course_name_entry.insert(0, "移动应用开发")
        
        # 班级名称
        ttk.Label(config_frame, text="班级名称:").pack(anchor=tk.W)
        self.class_name_entry = ttk.Entry(config_frame)
        self.class_name_entry.pack(fill=tk.X, pady=2)
        self.class_name_entry.insert(0, "软件23+6")
        
        # 课程目标权重配置
        ttk.Label(config_frame, text="课程目标权重配置:").pack(anchor=tk.W, pady=(10, 5))
        
        self.weight_entries = {}
        for obj in ['课程目标1', '课程目标2', '课程目标3', '课程目标4']:
            frame = ttk.Frame(config_frame)
            frame.pack(fill=tk.X, pady=2)
            ttk.Label(frame, text=f"{obj}:").pack(side=tk.LEFT)
            entry = ttk.Entry(frame, width=10)
            entry.pack(side=tk.RIGHT)
            entry.insert(0, str(self.objective_weights[obj]))
            self.weight_entries[obj] = entry
        
        # 保存配置按钮
        ttk.Button(config_frame, text="保存配置", 
                 command=self.save_config).pack(fill=tk.X, pady=10)
        
        # 计算按钮
        ttk.Button(config_frame, text="计算达成度", 
                 command=self.calculate_achievement).pack(fill=tk.X, pady=5)
        
        # 右侧面板 - 可视化和报告
        right_panel = ttk.LabelFrame(main_frame, text="可视化与报告", padding="10")
        right_panel.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5)
        right_panel.columnconfigure(0, weight=1)
        right_panel.rowconfigure(0, weight=1)
        
        # 结果显示区域
        result_frame = ttk.LabelFrame(right_panel, text="计算结果", padding="10")
        result_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 结果文本框
        self.result_text = scrolledtext.ScrolledText(result_frame, height=15, width=60)
        self.result_text.pack(fill=tk.BOTH, expand=True)
        
        # 可视化按钮
        viz_frame = ttk.Frame(right_panel)
        viz_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(viz_frame, text="生成柱状图", 
                 command=self.generate_bar_chart).pack(side=tk.LEFT, padx=2)
        ttk.Button(viz_frame, text="生成雷达图", 
                 command=self.generate_radar_chart).pack(side=tk.LEFT, padx=2)
        ttk.Button(viz_frame, text="生成分布图", 
                 command=self.generate_box_plot).pack(side=tk.LEFT, padx=2)
        
        # 报告生成按钮
        report_frame = ttk.Frame(right_panel)
        report_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(report_frame, text="生成Markdown报告", 
                 command=self.generate_markdown_report).pack(side=tk.LEFT, padx=2)
        ttk.Button(report_frame, text="生成Word报告", 
                 command=self.generate_word_report).pack(side=tk.LEFT, padx=2)
        ttk.Button(report_frame, text="打印报告", 
                 command=self.print_report).pack(side=tk.LEFT, padx=2)
        
        # 状态栏
        self.status_label = ttk.Label(main_frame, text="就绪", relief=tk.SUNKEN)
        self.status_label.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
    
    def import_excel(self):
        """导入Excel文件"""
        file_path = filedialog.askopenfilename(
            title="选择Excel成绩文件",
            filetypes=[("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")]
        )
        
        if file_path:
            try:
                self.excel_file_path = file_path
                self.excel_label.config(text=os.path.basename(file_path), foreground="black")
                self.update_status(f"已加载Excel文件: {os.path.basename(file_path)}")
                
                # 尝试解析Excel文件
                df = pd.read_excel(file_path)
                self.students = self.parse_excel_data(df)
                self.update_status(f"成功解析 {len(self.students)} 名学生数据")
                
            except Exception as e:
                messagebox.showerror("错误", f"导入Excel文件失败: {str(e)}")
                self.update_status("导入失败")
    
    def import_word(self):
        """导入Word文件"""
        file_path = filedialog.askopenfilename(
            title="选择Word大纲文件",
            filetypes=[("Word文件", "*.docx *.doc"), ("所有文件", "*.*")]
        )
        
        if file_path:
            try:
                self.syllabus_file_path = file_path
                self.word_label.config(text=os.path.basename(file_path), foreground="black")
                self.update_status(f"已加载Word文件: {os.path.basename(file_path)}")
                
                # 解析Word文件
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                self.course_config['syllabus'] = text
                self.update_status("成功解析大纲文件")
                
            except Exception as e:
                messagebox.showerror("错误", f"导入Word文件失败: {str(e)}")
                self.update_status("导入失败")
    
    def parse_excel_data(self, df):
        """解析Excel数据结构"""
        # 找到表头行
        header_row = -1
        for i, row in df.iterrows():
            if '课程目标' in str(row.iloc[0]):
                header_row = i
                break
        
        if header_row == -1:
            raise ValueError("未找到课程目标表头")
        
        # 提取数据行
        data_rows = []
        for i in range(header_row + 3, len(df)):
            row_data = []
            for col in df.columns:
                val = df.loc[i, col]
                if str(val) != 'nan':
                    row_data.append(val)
                else:
                    row_data.append('')
            if any(row_data):
                data_rows.append(row_data)
        
        # 提取学生数据
        students = []
        for row in data_rows:
            if len(row) >= 11:
                try:
                    student = {
                        '学号': str(row[0]),
                        '姓名': str(row[1]),
                        '目标1得分': float(row[2]) if row[2] else 0,
                        '目标1达成度': float(row[3]) if row[3] else 0,
                        '目标2得分': float(row[4]) if row[4] else 0,
                        '目标2达成度': float(row[5]) if row[5] else 0,
                        '目标3得分': float(row[6]) if row[6] else 0,
                        '目标3达成度': float(row[7]) if row[7] else 0,
                        '目标4得分': float(row[8]) if row[8] else 0,
                        '目标4达成度': float(row[9]) if row[9] else 0,
                        '总评得分': float(row[10]) if row[10] else 0
                    }
                    students.append(student)
                except Exception as e:
                    print(f"解析学生数据时出错: {e}")
                    continue
        
        return students
    
    def save_config(self):
        """保存配置"""
        try:
            # 更新权重
            for obj in ['课程目标1', '课程目标2', '课程目标3', '课程目标4']:
                try:
                    weight = float(self.weight_entries[obj].get())
                    self.objective_weights[obj] = weight
                except ValueError:
                    messagebox.showerror("错误", f"{obj}的权重必须是数字")
                    return
            
            # 验证权重总和
            total_weight = sum(self.objective_weights.values())
            if abs(total_weight - 1.0) > 0.01:
                messagebox.showwarning("警告", f"权重总和为{total_weight:.2f}，建议为1.0")
            
            # 保存课程信息
            self.course_config['course_name'] = self.course_name_entry.get()
            self.course_config['class_name'] = self.class_name_entry.get()
            self.course_config['objective_weights'] = self.objective_weights
            
            messagebox.showinfo("成功", "配置已保存")
            self.update_status("配置已保存")
            
        except Exception as e:
            messagebox.showerror("错误", f"保存配置失败: {str(e)}")
    
    def calculate_achievement(self):
        """计算达成度"""
        if not self.students:
            messagebox.showwarning("警告", "请先导入Excel文件")
            return
        
        try:
            # 计算班级平均达成度
            self.avg_achievements = self.calculate_class_average(self.students)
            
            # 计算加权达成度
            self.weighted_achievement = self.calculate_weighted_achievement(self.avg_achievements)
            
            # 显示结果
            self.display_results()
            
            self.update_status("计算完成")
            
        except Exception as e:
            messagebox.showerror("错误", f"计算达成度失败: {str(e)}")
            self.update_status("计算失败")
    
    def calculate_class_average(self, students):
        """计算班级平均达成度"""
        if not students:
            return {}
        
        avg_achievements = {
            '课程目标1': np.mean([s['目标1达成度'] for s in students]),
            '课程目标2': np.mean([s['目标2达成度'] for s in students]),
            '课程目标3': np.mean([s['目标3达成度'] for s in students]),
            '课程目标4': np.mean([s['目标4达成度'] for s in students]),
            '总评': np.mean([s['总评得分'] for s in students]) / 100
        }
        
        return avg_achievements
    
    def calculate_weighted_achievement(self, avg_achievements):
        """计算加权达成度"""
        weighted_achievement = 0
        for obj, weight in self.objective_weights.items():
            weighted_achievement += avg_achievements.get(obj, 0) * weight
        return weighted_achievement
    
    def display_results(self):
        """显示计算结果"""
        self.result_text.delete(1.0, tk.END)
        
        course_name = self.course_config.get('course_name', '移动应用开发')
        class_name = self.course_config.get('class_name', '软件23+6')
        
        result = f"""
{'='*60}
{class_name}《{course_name}》课程达成度计算结果
{'='*60}

一、班级平均达成度

课程目标1: {self.avg_achievements.get('课程目标1', 0):.4f} (权重: {self.objective_weights['课程目标1']:.2f})
课程目标2: {self.avg_achievements.get('课程目标2', 0):.4f} (权重: {self.objective_weights['课程目标2']:.2f})
课程目标3: {self.avg_achievements.get('课程目标3', 0):.4f} (权重: {self.objective_weights['课程目标3']:.2f})
课程目标4: {self.avg_achievements.get('课程目标4', 0):.4f} (权重: {self.objective_weights['课程目标4']:.2f})

加权总达成度: {self.weighted_achievement:.4f}

二、学生统计

学生总数: {len(self.students)}
平均总评: {self.avg_achievements.get('总评', 0):.4f}

三、达成度等级

"""
        
        if self.weighted_achievement >= 0.85:
            result += "优秀 (≥0.85)\n"
        elif self.weighted_achievement >= 0.70:
            result += "良好 (0.70-0.84)\n"
        elif self.weighted_achievement >= 0.60:
            result += "中等 (0.60-0.69)\n"
        else:
            result += "未达成 (<0.60)\n"
        
        result += f"{'='*60}\n"
        
        self.result_text.insert(tk.END, result)
    
    def generate_bar_chart(self):
        """生成柱状图"""
        if not self.avg_achievements:
            messagebox.showwarning("警告", "请先计算达成度")
            return
        
        try:
            objectives = list(self.avg_achievements.keys())[:-1]
            achievements = [self.avg_achievements[obj] for obj in objectives]
            
            plt.figure(figsize=(10, 6))
            bars = plt.bar(objectives, achievements, 
                        color=['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4'])
            plt.title('课程目标达成度', fontsize=14, fontweight='bold')
            plt.ylabel('达成度', fontsize=12)
            plt.ylim(0, 1)
            
            for bar, val in zip(bars, achievements):
                plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01,
                        f'{val:.4f}', ha='center', va='bottom', fontsize=10)
            
            plt.grid(axis='y', alpha=0.3)
            plt.tight_layout()
            
            output_path = os.path.join(self.output_dir, '课程目标达成度柱状图.png')
            plt.savefig(output_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            messagebox.showinfo("成功", f"柱状图已保存到: {output_path}")
            self.update_status(f"已生成柱状图: {output_path}")
            
        except Exception as e:
            messagebox.showerror("错误", f"生成柱状图失败: {str(e)}")
    
    def generate_radar_chart(self):
        """生成雷达图"""
        if not self.avg_achievements:
            messagebox.showwarning("警告", "请先计算达成度")
            return
        
        try:
            from math import pi
            
            objectives = list(self.avg_achievements.keys())[:-1]
            achievements = [self.avg_achievements[obj] for obj in objectives]
            
            plt.figure(figsize=(8, 8))
            ax = plt.subplot(111, polar=True)
            
            categories = objectives
            values = achievements
            values += values[:1]
            
            angles = [n / float(len(categories)) * 2 * pi for n in range(len(categories))]
            angles += angles[:1]
            
            ax.plot(angles, values, 'o-', linewidth=2, label='达成度', color='#FF6B6B')
            ax.fill(angles, values, alpha=0.25, color='#FF6B6B')
            ax.set_thetagrids([a * 180/pi for a in angles[:-1]], categories)
            ax.set_ylim(0, 1)
            plt.title('课程目标达成度雷达图', fontsize=14, fontweight='bold', pad=20)
            plt.legend(loc='upper right', bbox_to_anchor=(0.1, 0.1))
            plt.tight_layout()
            
            output_path = os.path.join(self.output_dir, '课程目标达成度雷达图.png')
            plt.savefig(output_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            messagebox.showinfo("成功", f"雷达图已保存到: {output_path}")
            self.update_status(f"已生成雷达图: {output_path}")
            
        except Exception as e:
            messagebox.showerror("错误", f"生成雷达图失败: {str(e)}")
    
    def generate_box_plot(self):
        """生成箱线图"""
        if not self.students:
            messagebox.showwarning("警告", "请先导入数据")
            return
        
        try:
            student_achievements = []
            for student in self.students:
                student_achievements.append([
                    student['目标1达成度'],
                    student['目标2达成度'],
                    student['目标3达成度'],
                    student['目标4达成度']
                ])
            
            student_achievements = np.array(student_achievements)
            
            plt.figure(figsize=(12, 6))
            plt.boxplot(student_achievements, 
                      labels=['课程目标1', '课程目标2', '课程目标3', '课程目标4'])
            plt.title('学生个体达成度分布', fontsize=14, fontweight='bold')
            plt.ylabel('达成度', fontsize=12)
            plt.grid(axis='y', alpha=0.3)
            plt.tight_layout()
            
            output_path = os.path.join(self.output_dir, '学生个体达成度分布图.png')
            plt.savefig(output_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            messagebox.showinfo("成功", f"分布图已保存到: {output_path}")
            self.update_status(f"已生成分布图: {output_path}")
            
        except Exception as e:
            messagebox.showerror("错误", f"生成分布图失败: {str(e)}")
    
    def generate_markdown_report(self):
        """生成Markdown报告"""
        if not self.avg_achievements:
            messagebox.showwarning("警告", "请先计算达成度")
            return
        
        try:
            course_name = self.course_config.get('course_name', '移动应用开发')
            class_name = self.course_config.get('class_name', '软件23+6')
            
            report_content = f"""# {class_name}《{course_name}》课程达成度报告

**生成时间：** {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}

---

## 一、课程目标达成情况

### 1. 班级平均达成度

| 课程目标 | 达成度 | 权重 | 加权贡献 |
|---------|-------|------|---------|
| 课程目标1 | {self.avg_achievements.get('课程目标1', 0):.4f} | {self.objective_weights['课程目标1']:.2f} | {self.avg_achievements.get('课程目标1', 0) * self.objective_weights['课程目标1']:.4f} |
| 课程目标2 | {self.avg_achievements.get('课程目标2', 0):.4f} | {self.objective_weights['课程目标2']:.2f} | {self.avg_achievements.get('课程目标2', 0) * self.objective_weights['课程目标2']:.4f} |
| 课程目标3 | {self.avg_achievements.get('课程目标3', 0):.4f} | {self.objective_weights['课程目标3']:.2f} | {self.avg_achievements.get('课程目标3', 0) * self.objective_weights['课程目标3']:.4f} |
| 课程目标4 | {self.avg_achievements.get('课程目标4', 0):.4f} | {self.objective_weights['课程目标4']:.2f} | {self.avg_achievements.get('课程目标4', 0) * self.objective_weights['课程目标4']:.4f} |
| **加权总达成度** | **{self.weighted_achievement:.4f}** | **1.00** | **{self.weighted_achievement:.4f}** |

### 2. 学生个体达成情况

共有 **{len(self.students)}** 名学生参与评价。

#### 学生达成度统计

| 统计指标 | 课程目标1 | 课程目标2 | 课程目标3 | 课程目标4 |
|---------|----------|----------|----------|----------|
| 平均值 | {np.mean([s['目标1达成度'] for s in self.students]):.4f} | {np.mean([s['目标2达成度'] for s in self.students]):.4f} | {np.mean([s['目标3达成度'] for s in self.students]):.4f} | {np.mean([s['目标4达成度'] for s in self.students]):.4f} |
| 最大值 | {np.max([s['目标1达成度'] for s in self.students]):.4f} | {np.max([s['目标2达成度'] for s in self.students]):.4f} | {np.max([s['目标3达成度'] for s in self.students]):.4f} | {np.max([s['目标4达成度'] for s in self.students]):.4f} |
| 最小值 | {np.min([s['目标1达成度'] for s in self.students]):.4f} | {np.min([s['目标2达成度'] for s in self.students]):.4f} | {np.min([s['目标3达成度'] for s in self.students]):.4f} | {np.min([s['目标4达成度'] for s in self.students]):.4f} |
| 标准差 | {np.std([s['目标1达成度'] for s in self.students]):.4f} | {np.std([s['目标2达成度'] for s in self.students]):.4f} | {np.std([s['目标3达成度'] for s in self.students]):.4f} | {np.std([s['目标4达成度'] for s in self.students]):.4f} |

---

## 二、达成度分析

### 1. 定量评价情况分析

#### 课程目标1分析
**达成度：** {self.avg_achievements.get('课程目标1', 0):.4f}

从达成度结果可以看出，学生在课程目标1方面表现{("良好" if self.avg_achievements.get('课程目标1', 0) >= 0.7 else "一般")}。

#### 课程目标2分析
**达成度：** {self.avg_achievements.get('课程目标2', 0):.4f}

从达成度结果可以看出，学生在课程目标2方面表现{("良好" if self.avg_achievements.get('课程目标2', 0) >= 0.7 else "一般")}。

#### 课程目标3分析
**达成度：** {self.avg_achievements.get('课程目标3', 0):.4f}

从达成度结果可以看出，学生在课程目标3方面表现{("良好" if self.avg_achievements.get('课程目标3', 0) >= 0.7 else "一般")}。

#### 课程目标4分析
**达成度：** {self.avg_achievements.get('课程目标4', 0):.4f}

从达成度结果可以看出，学生在课程目标4方面表现{("良好" if self.avg_achievements.get('课程目标4', 0) >= 0.7 else "一般")}。

---

## 三、结论

通过本次课程达成度评价，我们可以看到：

1. **整体表现**：学生在{course_name}课程的学习中取得了一定的成果，加权总达成度为{self.weighted_achievement:.4f}。

2. **达成度等级**：
"""
            
            if self.weighted_achievement >= 0.85:
                report_content += "优秀 (≥0.85)\n"
            elif self.weighted_achievement >= 0.70:
                report_content += "良好 (0.70-0.84)\n"
            elif self.weighted_achievement >= 0.60:
                report_content += "中等 (0.60-0.69)\n"
            else:
                report_content += "未达成 (<0.60)\n"
            
            report_content += """
3. **改进方向**：通过持续的教学改进，我们相信学生的能力将得到进一步提升。

---

**报告生成完成**
"""
            
            output_path = os.path.join(self.output_dir, f'{class_name}《{course_name}》课程达成度报告.md')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            
            messagebox.showinfo("成功", f"报告已保存到: {output_path}")
            self.update_status(f"已生成Markdown报告: {output_path}")
            
        except Exception as e:
            messagebox.showerror("错误", f"生成报告失败: {str(e)}")
    
    def generate_word_report(self):
        """生成Word报告"""
        if not self.avg_achievements:
            messagebox.showwarning("警告", "请先计算达成度")
            return
        
        try:
            course_name = self.course_config.get('course_name', '移动应用开发')
            class_name = self.course_config.get('class_name', '软件23+6')
            
            doc = Document()
            
            # 标题
            title = doc.add_heading(f'{class_name}《{course_name}》课程达成度报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 生成时间
            doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
            doc.add_paragraph()
            
            # 课程目标达成情况
            doc.add_heading('一、课程目标达成情况', level=1)
            doc.add_heading('1. 班级平均达成度', level=2)
            
            # 创建表格
            table = doc.add_table(rows=6, cols=4)
            table.style = 'Table Grid'
            
            # 表头
            headers = ['课程目标', '达成度', '权重', '加权贡献']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            # 数据行
            objectives = ['课程目标1', '课程目标2', '课程目标3', '课程目标4']
            for i, obj in enumerate(objectives):
                table.rows[i+1].cells[0].text = obj
                table.rows[i+1].cells[1].text = f"{self.avg_achievements.get(obj, 0):.4f}"
                table.rows[i+1].cells[2].text = f"{self.objective_weights[obj]:.2f}"
                table.rows[i+1].cells[3].text = f"{self.avg_achievements.get(obj, 0) * self.objective_weights[obj]:.4f}"
            
            # 总计行
            table.rows[5].cells[0].text = '加权总达成度'
            table.rows[5].cells[1].text = f"{self.weighted_achievement:.4f}"
            table.rows[5].cells[2].text = '1.00'
            table.rows[5].cells[3].text = f"{self.weighted_achievement:.4f}"
            
            # 保存文档
            output_path = os.path.join(self.output_dir, f'{class_name}《{course_name}》课程达成度报告.docx')
            doc.save(output_path)
            
            messagebox.showinfo("成功", f"Word报告已保存到: {output_path}")
            self.update_status(f"已生成Word报告: {output_path}")
            
        except Exception as e:
            messagebox.showerror("错误", f"生成Word报告失败: {str(e)}")
    
    def print_report(self):
        """打印报告"""
        if not self.avg_achievements:
            messagebox.showwarning("警告", "请先计算达成度")
            return
        
        try:
            # 先生成Markdown报告
            self.generate_markdown_report()
            
            # 打开生成的报告文件
            course_name = self.course_config.get('course_name', '移动应用开发')
            class_name = self.course_config.get('class_name', '软件23+6')
            report_path = os.path.join(self.output_dir, f'{class_name}《{course_name}》课程达成度报告.md')
            
            if os.path.exists(report_path):
                os.startfile(report_path)
                messagebox.showinfo("提示", f"已打开报告文件，请使用浏览器的打印功能进行打印")
                self.update_status("已打开报告文件")
            else:
                messagebox.showerror("错误", "报告文件不存在")
                
        except Exception as e:
            messagebox.showerror("错误", f"打印报告失败: {str(e)}")
    
    def update_status(self, message):
        """更新状态栏"""
        self.status_label.config(text=message)


def main():
    root = tk.Tk()
    app = CourseAchievementGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()